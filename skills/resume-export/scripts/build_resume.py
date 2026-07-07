#!/usr/bin/env python3
"""
Rebound — generic, JSON-driven resume renderer.

Usage:
    python3 build_resume.py <config.json>

Produces a clean, single-column, ATS-safe .docx (no tables / text boxes / columns).
The config is fully generic — nothing candidate-specific is hard-coded here; the
/rebound:tailor command assembles the config from the user's BASE resume + the JD.

Config schema (all fields optional unless noted):
{
  "out": "/abs/path/Output.docx",            # required
  "name": "JANE DOE",                          # required
  "title": "Senior Software Engineer  |  ...", # tagline under the name
  "contact": "+1 ... | jane@x.com | linkedin.com/in/jane | City, ST",
  "accent": "1F3A5F",                          # hex heading color (default deep navy)
  "summary": "3-4 sentence summary ...",
  "sections": [
    {"heading": "Work Experience", "type": "experience",
     "roles": [{"title": "Senior Engineer — ACME", "meta": "City, ST | 2022 – Present",
                "bullets": ["**Led** ...", "..."]}]},
    {"heading": "Skills", "type": "skills",
     "skills": [{"label": "Languages", "items": "Java, Go, Python"}]},
    {"heading": "Education", "type": "roles_only",
     "roles": [{"title": "B.S. CS — University", "meta": "GPA 3.9 | 2016"}]},
    {"heading": "Certifications", "type": "bullets", "bullets": ["**AWS** ...", "..."]}
  ]
}

Inline markup in any text/bullet: **bold**, *italic*.
"""
import sys, json, os

# Allow a bundled python-docx (installed by the export skill) to be found.
for extra in (os.environ.get("REBOUND_PYLIBS"), "/tmp/pylibs",
              os.path.expanduser("~/.rebound/pylibs")):
    if extra and os.path.isdir(extra):
        sys.path.insert(0, extra)

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x20, 0x20, 0x20)
GREY = RGBColor(0x55, 0x55, 0x55)
FONT = "Calibri"


def hexcolor(h, default=(0x1F, 0x3A, 0x5F)):
    try:
        h = (h or "").lstrip("#")
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
    except Exception:
        return RGBColor(*default)


def build(cfg):
    accent = hexcolor(cfg.get("accent"))
    accent_hex = (cfg.get("accent") or "1F3A5F").lstrip("#")
    doc = Document()

    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(10)
    st.font.color.rgb = DARK
    pf = st.paragraph_format
    pf.space_after = Pt(0); pf.space_before = Pt(0); pf.line_spacing = 1.0

    for s in doc.sections:
        s.top_margin = Inches(0.5); s.bottom_margin = Inches(0.5)
        s.left_margin = Inches(0.6); s.right_margin = Inches(0.6)
    PAGE_W = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin

    def space(p, before=0, after=0, line=1.0):
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing = line

    def run(p, text, size=10, bold=False, italic=False, color=DARK):
        r = p.add_run(text)
        r.font.name = FONT; r.font.size = Pt(size)
        r.bold = bold; r.italic = italic; r.font.color.rgb = color
        return r

    def rich(p, text, size=10, color=DARK):
        """Render **bold** and *italic* inline markers into runs."""
        for bi, bseg in enumerate((text or "").split("**")):
            if bseg == "":
                continue
            if bi % 2 == 1:
                run(p, bseg, size=size, bold=True, color=color)
            else:
                for ii, iseg in enumerate(bseg.split("*")):
                    if iseg == "":
                        continue
                    run(p, iseg, size=size, italic=(ii % 2 == 1), color=color)

    def rule(p):
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), accent_hex)
        pbdr.append(bottom); pPr.append(pbdr)

    def heading(text):
        p = doc.add_paragraph(); space(p, before=7, after=3)
        run(p, text.upper(), size=11, bold=True, color=accent); rule(p)

    def role_header(title, meta):
        p = doc.add_paragraph(); space(p, before=5, after=1)
        p.paragraph_format.tab_stops.add_tab_stop(PAGE_W, WD_TAB_ALIGNMENT.RIGHT)
        run(p, title, size=10.5, bold=True, color=DARK)
        if meta:
            run(p, "\t" + meta, size=9.5, italic=True, color=GREY)

    def bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.first_line_indent = Inches(-0.13)
        space(p, before=0, after=2, line=1.02)
        rich(p, text)

    # ---- header ----
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; space(p, after=1)
    run(p, cfg.get("name", ""), size=22, bold=True, color=accent)
    if cfg.get("title"):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; space(p, after=2)
        run(p, cfg["title"], size=10.5, bold=True, color=GREY)
    if cfg.get("contact"):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; space(p, after=4)
        run(p, cfg["contact"], size=9.5, color=GREY)

    # ---- summary ----
    if cfg.get("summary"):
        heading("Summary")
        p = doc.add_paragraph(); space(p, before=2, after=2, line=1.05)
        rich(p, cfg["summary"])

    # ---- sections ----
    for sec in cfg.get("sections", []):
        heading(sec.get("heading", ""))
        stype = sec.get("type", "bullets")
        if stype == "experience":
            for role in sec.get("roles", []):
                role_header(role.get("title", ""), role.get("meta", ""))
                for b in role.get("bullets", []):
                    bullet(b)
        elif stype == "roles_only":
            for role in sec.get("roles", []):
                role_header(role.get("title", ""), role.get("meta", ""))
        elif stype == "skills":
            for sk in sec.get("skills", []):
                p = doc.add_paragraph(); space(p, before=1, after=1, line=1.05)
                run(p, sk.get("label", "") + ":  ", size=10, bold=True, color=accent)
                run(p, sk.get("items", ""), size=10, color=DARK)
        else:  # bullets
            for b in sec.get("bullets", []):
                bullet(b)

    out = cfg["out"]
    os.makedirs(os.path.dirname(os.path.abspath(out)), exist_ok=True)
    doc.save(out)
    return out


def main():
    if len(sys.argv) < 2:
        print("usage: python3 build_resume.py <config.json>", file=sys.stderr)
        sys.exit(2)
    with open(sys.argv[1], "r", encoding="utf-8") as f:
        cfg = json.load(f)
    out = build(cfg)
    print("Saved:", out)


if __name__ == "__main__":
    main()
